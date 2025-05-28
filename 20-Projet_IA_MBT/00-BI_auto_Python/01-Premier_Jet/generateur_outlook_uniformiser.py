import gradio as gr
from docx import Document
import subprocess
import sys
import os
from docx2pdf import convert
from datetime import datetime
import win32com.client as win32
import pythoncom
import re

# Dictionnaire des intervenants
intervenants = {
    "Mounir BOUGOUFFA": "mbougouffa@mbt-consulting.com",
    "Yannis CHAZOT": "ychazot@mbt-consulting.com",
}

def get_mail_intervenant(nom_intervenant):
    return intervenants.get(nom_intervenant, "")

def validate_inputs(societe, nom_contact, mail_contact, duree_inter, date_deb, date_fin, obj_presta, contenu_intervention, num_mission):
    errors = []
    if not re.match(r"[^@]+@[^@]+\.[^@]+", mail_contact):
        errors.append("Adresse email invalide.")
    try:
        datetime.strptime(date_deb, "%d/%m/%Y")
        datetime.strptime(date_fin, "%d/%m/%Y")
    except ValueError:
        errors.append("Format de date invalide. Utilisez le format JJ/MM/AAAA.")
    return errors

def replace_text_in_runs(runs, placeholder, replacement):
    for run in runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)

def print_document(file_path):
    try:
        subprocess.run([
            'C:\\Program Files (x86)\\SumatraPDF\\SumatraPDF.exe',
            '-silent',
            file_path
        ], check=True)
    except Exception as e:
        print(f"Erreur d'impression : {e}")

def ouvrir_email_outlook(mail_contact, intervenant, societe, output_pdf):
    pythoncom.CoInitialize()
    cc_emails = ";".join([email for name, email in intervenants.items() if name != intervenant])
    subject = f"MBT/{societe} - Bon d'intervention"
    body = "Bonjour,\n\nVeuillez trouver ci-joint le bon d'intervention à nous retourner signé.\n\n"
    outlook = win32.Dispatch('Outlook.Application')
    mail = outlook.CreateItem(0)
    mail.To = mail_contact
    mail.CC = cc_emails
    mail.Subject = subject
    mail.Body = body
    mail.Attachments.Add(os.path.abspath(output_pdf))
    mail.Display()
    pythoncom.CoUninitialize()

def generate_document(template, intervenant, societe, nom_contact, mail_contact, duree_inter, date_deb, date_fin, obj_presta, contenu_intervention, num_mission):
    if not num_mission:
        num_mission = "PRXXXXX-XX"

    errors = validate_inputs(societe, nom_contact, mail_contact, duree_inter, date_deb, date_fin, obj_presta, contenu_intervention, num_mission)
    if errors:
        return None, "\n".join(errors)

    date = datetime.today().strftime("%d/%m/%Y")
    mail_intervenant = get_mail_intervenant(intervenant)
    doc = Document(template)

    remplacements = {
        '[DATE]': date, '[INTERVENANT]': intervenant, '[MAIL_INTERVENANT]': mail_intervenant,
        '[SOCIETE]': societe, '[NOM_CONTACT]': nom_contact, '[MAIL_CONTACT]': mail_contact,
        '[DUREE_INTER]': duree_inter, '[DATE_DEB]': date_deb, '[DATE_FIN]': date_fin,
        '[OBJ_PRESTA]': obj_presta, '[CONTENU_INTERVENTION]': contenu_intervention,
        '[NUM_MISSION]': num_mission
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in remplacements.items():
            replace_text_in_runs(paragraph.runs, placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in remplacements.items():
                        replace_text_in_runs(paragraph.runs, placeholder, value)

    for section in doc.sections:
        for part in [section.header, section.footer]:
            for paragraph in part.paragraphs:
                for placeholder, value in remplacements.items():
                    replace_text_in_runs(paragraph.runs, placeholder, value)

    output_docx = f"BI_{societe}_{datetime.today().strftime('%Y%m%d')}.docx"
    output_pdf = output_docx.replace(".docx", ".pdf")
    doc.save(output_docx)
    pythoncom.CoInitialize()
    convert(output_docx, output_pdf)
    pythoncom.CoUninitialize()
    os.remove(output_docx)

    try:
        print_document(output_pdf)
        ouvrir_email_outlook(mail_contact, intervenant, societe, output_pdf)
    except Exception as e:
        return None, f"Erreur lors de l'ouverture du mail ou de l'impression : {e}"

    return output_pdf, "Document généré et mail prêt à être envoyé."

# Interface Gradio
with gr.Blocks() as iface:
    template_input = gr.Dropdown(["template_bon-intervention.docx"], label="Template")
    num_mission_input = gr.Textbox(label="Numéro de Mission")
    intervenant_input = gr.Dropdown(choices=list(intervenants.keys()), label="Intervenant")
    societe_input = gr.Textbox(label="Société")
    nom_contact_input = gr.Textbox(label="Nom Contact")
    mail_contact_input = gr.Textbox(label="Mail Contact")
    duree_input = gr.Textbox(label="Durée Intervention")
    date_deb_input = gr.Textbox(label="Date Début (JJ/MM/AAAA)")
    date_fin_input = gr.Textbox(label="Date Fin (JJ/MM/AAAA)")
    obj_presta_input = gr.Textbox(label="Objectif Prestation")
    contenu_input = gr.Textbox(label="Contenu Intervention")
    output_file = gr.File(label="Bon d'intervention (PDF)")
    log_output = gr.Textbox(label="Log", interactive=False)

    bouton = gr.Button("Générer le PDF et le mail")
    bouton.click(
        fn=generate_document,
        inputs=[
            template_input, intervenant_input, societe_input, nom_contact_input,
            mail_contact_input, duree_input, date_deb_input, date_fin_input,
            obj_presta_input, contenu_input, num_mission_input
        ],
        outputs=[output_file, log_output]
    )

    iface.launch()
