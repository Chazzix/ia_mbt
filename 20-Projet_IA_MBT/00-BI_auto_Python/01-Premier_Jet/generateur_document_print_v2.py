# -*- coding: utf-8 -*-
import gradio as gr
from docx import Document
import subprocess
import sys
import os
import platform
from docx2pdf import convert
from datetime import datetime

# Verifie et installe les bibliotheques necessaires
def check_install_libraries():
    required_libraries = ["gradio", "python-docx", "docx2pdf"]
    for lib in required_libraries:
        try:
            __import__(lib)
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", lib])

check_install_libraries()

def replace_text_in_runs(runs, placeholder, replacement):
    for run in runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)

# Dictionnaire des intervenants présents chez MBT
intervenants = {
    "Mounir BOUGOUFFA": "mbougouffa@mbt-consulting.com",
    "Yannis CHAZOT": "ychazot@mbt-consulting.com",
}

# Ajout de l'email intervenant auto
def get_mail_intervenant(nom_intervenant):
    return intervenants.get(nom_intervenant, "")

# Impression du documents
def print_document(file_path):
    try:
        if not isinstance(file_path, (str, bytes, os.PathLike)):
            raise ValueError(f"Chemin de fichier invalide : {file_path}")
        
        subprocess.run([
            'D:\\10-Tools\\08-SumatraPDF\\SumatraPDF\\SumatraPDF.exe',
            '-silent',
            file_path
        ], check=True)
    except subprocess.CalledProcessError as e:
        print(f"Erreur d'impression (code {e.returncode}) : {e}")
    except Exception as e:
        print(f"Erreur inattendue : {e}")

def generate_document(template, intervenant, societe, nom_contact, mail_contact, duree_inter, date_deb, date_fin, obj_presta, contenu_intervention):
    date = datetime.today().strftime("%d/%m/%Y")
    mail_intervenant = get_mail_intervenant(intervenant)
    doc = Document(template)

    # Remplacement dans les paragraphes
    for paragraph in doc.paragraphs:
        for placeholder, value in {
            '[DATE]': date, '[INTERVENANT]': intervenant, '[MAIL_INTERVENANT]': mail_intervenant,
            '[SOCIETE]': societe, '[NOM_CONTACT]': nom_contact, '[MAIL_CONTACT]': mail_contact,
            '[DUREE_INTER]': duree_inter, '[DATE_DEB]': date_deb, '[DATE_FIN]': date_fin,
            '[OBJ_PRESTA]': obj_presta, '[CONTENU_INTERVENTION]': contenu_intervention
        }.items():
            replace_text_in_runs(paragraph.runs, placeholder, value)

    # Remplacement dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in {
                        '[DATE]': date, '[INTERVENANT]': intervenant, '[MAIL_INTERVENANT]': mail_intervenant,
                        '[SOCIETE]': societe, '[NOM_CONTACT]': nom_contact, '[MAIL_CONTACT]': mail_contact,
                        '[DUREE_INTER]': duree_inter, '[DATE_DEB]': date_deb, '[DATE_FIN]': date_fin,
                        '[OBJ_PRESTA]': obj_presta, '[CONTENU_INTERVENTION]': contenu_intervention
                    }.items():
                        replace_text_in_runs(paragraph.runs, placeholder, value)

    # En-t�tes et pieds de page
    for section in doc.sections:
        for part in [section.header, section.footer]:
            for paragraph in part.paragraphs:
                for placeholder, value in {
                    '[DATE]': date, '[INTERVENANT]': intervenant, '[MAIL_INTERVENANT]': mail_intervenant,
                    '[SOCIETE]': societe, '[NOM_CONTACT]': nom_contact, '[MAIL_CONTACT]': mail_contact,
                    '[DUREE_INTER]': duree_inter, '[DATE_DEB]': date_deb, '[DATE_FIN]': date_fin,
                    '[OBJ_PRESTA]': obj_presta, '[CONTENU_INTERVENTION]': contenu_intervention
                }.items():
                    replace_text_in_runs(paragraph.runs, placeholder, value)

    # Sauvegarde et conversion
    output_docx = f"BI_{societe}_{datetime.today().strftime('%Y%m%d')}.docx"
    output_pdf = output_docx.replace(".docx", ".pdf")
    doc.save(output_docx)
    convert(output_docx, output_pdf)
    os.remove(output_docx)

    # Impression
    print_document(output_pdf)

    return output_pdf

# Interface Gradio
with gr.Blocks() as iface:
    template_input = gr.Dropdown(["template_bon-intervention.docx"], label="Template")
    intervenant_input = gr.Dropdown(choices=list(intervenants.keys()), label="Intervenant")
    societe_input = gr.Textbox(label="Société")
    nom_contact_input = gr.Textbox(label="Nom Contact")
    mail_contact_input = gr.Textbox(label="Mail Contact")
    duree_input = gr.Textbox(label="Durée Intervention")
    date_deb_input = gr.Textbox(label="Date Début")
    date_fin_input = gr.Textbox(label="Date Fin")
    obj_presta_input = gr.Textbox(label="Objectif Prestation")
    contenu_input = gr.Textbox(label="Contenu Intervention")
    output_file = gr.File(label="Bon d'intervention (PDF)")

    bouton = gr.Button("Générer le PDF")

    bouton.click(
        fn=generate_document,
        inputs=[
            template_input,
            intervenant_input,
            societe_input,
            nom_contact_input,
            mail_contact_input,
            duree_input,
            date_deb_input,
            date_fin_input,
            obj_presta_input,
            contenu_input
        ],
        outputs=output_file
    )

iface.launch()