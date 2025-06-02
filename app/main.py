import gradio as gr
import psycopg2
from docx import Document
import os
from datetime import datetime
from dotenv import load_dotenv
import sys
import smtplib
from email.message import EmailMessage

load_dotenv()

DB_HOST = os.getenv("DB_HOST")
DB_NAME = os.getenv("DB_NAME")
DB_USER = os.getenv("DB_USER")
DB_PASS = os.getenv("DB_PASS")

def connect_db():
    return psycopg2.connect(
        host=DB_HOST,
        database=DB_NAME,
        user=DB_USER,
        password=DB_PASS
    )

def get_clients():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT societe FROM clients")
    clients = [row[0] for row in cur.fetchall()]
    cur.close()
    conn.close()
    return clients

def get_contacts(societe):
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT nom FROM contact WHERE client_id=(SELECT id FROM clients WHERE societe=%s)", (societe,))
    contacts = [row[0] for row in cur.fetchall()]
    cur.close()
    conn.close()
    return contacts

def get_intervenants():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT intervenant FROM intervenants")
    intervenants = [row[0] for row in cur.fetchall()]
    cur.close()
    conn.close()
    return intervenants

def get_mail_intervenant(nom_intervenant):
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT mail FROM intervenants WHERE intervenant = %s", (nom_intervenant,))
    mail = cur.fetchone()[0]
    cur.close()
    conn.close()
    return mail

def get_all_intervenant_emails(exclude_email=None):
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT mail FROM intervenants")
    mails = [row[0] for row in cur.fetchall()]
    cur.close()
    conn.close()
    if exclude_email:
        mails = [m for m in mails if m != exclude_email]
    return mails

def get_mail_contact(nom_contact):
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT mail FROM contact WHERE nom = %s", (nom_contact,))
    mail = cur.fetchone()[0]
    cur.close()
    conn.close()
    return mail

def add_client(nomclient):
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("INSERT INTO clients (societe) VALUES (%s)", (nomclient,))
    conn.commit()
    cur.close()
    conn.close()
    return f"Client '{nomclient}' ajouté avec succès."

def add_contact(societe, nom, prenom, mail, telephone):
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO contact (nom, prenom, mail, telephone, client_id)
        VALUES (%s, %s, %s, %s, (SELECT id FROM clients WHERE societe=%s))
    """, (nom, prenom, mail, telephone, societe))
    conn.commit()
    cur.close()
    conn.close()
    return f"Contact {nom} ajouté avec succès."

def get_bon_intervention():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM bon_intervention_view")
    data = cur.fetchall()
    cur.close()
    conn.close()
    return data

def replace_placeholders(text, replacements):
    for key, value in replacements.items():
        text = text.replace(key, value)
    return text

def generate_document(intervenant, societe, contact, duree_inter, date_deb, date_fin, obj_presta, contenu_intervention, num_mission, mail_intervenant):
    doc = Document("template_bon-intervention.docx")
    replacements = {
        "[INTERVENANT]": intervenant,
        "[MAIL_INTERVENANT]": mail_intervenant,
        "[SOCIETE]": societe,
        "[NOM_CONTACT]": contact,
        "[DUREE_INTER]": duree_inter,
        "[DATE_DEB]": date_deb,
        "[DATE_FIN]": date_fin,
        "[OBJ_PRESTA]": obj_presta,
        "[CONTENU_INTERVENTION]": contenu_intervention,
        "[NUM_MISSION]": num_mission,
        "[DATE]": datetime.today().strftime("%d/%m/%Y")
    }
    for p in doc.paragraphs:
        p.text = replace_placeholders(p.text, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_placeholders(cell.text, replacements)
        # p.text = p.text.replace("[INTERVENANT]", intervenant)\
                    #    .replace("[MAIL_INTERVENANT]", mail_intervenant)\
                    #    .replace("[SOCIETE]", societe)\
                    #    .replace("[NOM_CONTACT]", contact)\
                    #    .replace("[DUREE_INTER]", duree_inter)\
                    #    .replace("[DATE_DEB]", date_deb)\
                    #    .replace("[DATE_FIN]", date_fin)\
                    #    .replace("[OBJ_PRESTA]", obj_presta)\
                    #    .replace("[CONTENU_INTERVENTION]", contenu_intervention)\
                    #    .replace("[NUM_MISSION]", num_mission)\
                    #    .replace("[DATE]", datetime.today().strftime("%d/%m/%Y"))
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             cell.text = cell.text.replace("[INTERVENANT]", intervenant)\
    #                                 .replace("[MAIL_INTERVENANT]", mail_intervenant)\
    #                                 .replace("[SOCIETE]", societe)\
    #                                 .replace("[NOM_CONTACT]", contact)\
    #                                 .replace("[DUREE_INTER]", duree_inter)\
    #                                 .replace("[DATE_DEB]", date_deb)\
    #                                 .replace("[DATE_FIN]", date_fin)\
    #                                 .replace("[OBJ_PRESTA]", obj_presta)\
    #                                 .replace("[CONTENU_INTERVENTION]", contenu_intervention)\
    #                                 .replace("[NUM_MISSION]", num_mission)\
    #                                 .replace("[DATE]", datetime.today().strftime("%d/%m/%Y"))
    doc_path = f"BI_{societe.replace(' ', '_')}.docx"
    doc.save(doc_path)

    # Conversion en PDF avec LibreOffice
    os.system(f'libreoffice --headless --convert-to pdf "{doc_path}" --outdir .')
    pdf_path = doc_path.replace(".docx", ".pdf")

    # Suppression du fichier Word
    if os.path.exists(doc_path):
        os.remove(doc_path)

    # Ajoute cette conversion avant l'insertion SQL
    date_deb = datetime.strptime(date_deb, "%d/%m/%Y").date()
    date_fin = datetime.strptime(date_fin, "%d/%m/%Y").date()

    # Insertion en base
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO bon_intervention (
            intervenant_id, client_id, contact_id, duree_inter, date_deb, date_fin, obj_presta, contenu_intervention, num_mission
        )
        VALUES (
            (SELECT id FROM intervenants WHERE intervenant=%s),
            (SELECT id FROM clients WHERE societe=%s),
            (SELECT id FROM contact WHERE nom=%s),
            %s, %s, %s, %s, %s, %s
        )
    """, (intervenant, societe, contact, duree_inter, date_deb, date_fin, obj_presta, contenu_intervention, num_mission))
    conn.commit()
    cur.close()
    conn.close()

    return pdf_path

def prepare_outlook_email(mail_contact, mail_intervenant, pdf_path, societe):
    from email.message import EmailMessage

    cc_list = get_all_intervenant_emails(exclude_email=mail_intervenant)

    msg = EmailMessage()
    msg["Subject"] = f"MBT/{societe} - Bon d'intervention"
    msg["From"] = mail_intervenant
    msg["To"] = mail_contact
    if cc_list:
        msg["Cc"] = ", ".join(cc_list)
    msg.set_content("Bonjour,\n\nVeuillez trouver ci-joint le bon d'intervention.\n\n")

    with open(pdf_path, "rb") as f:
        file_data = f.read()
        file_name = os.path.basename(pdf_path)
        msg.add_attachment(file_data, maintype="application", subtype="pdf", filename=file_name)
    
    output_dir = "./emails"
    os.makedirs(output_dir, exist_ok=True)
    eml_path = os.path.join(output_dir, f"email_{file_name.replace('.pdf', '.eml')}")

    with open(eml_path, "wb") as f:
        f.write(msg.as_bytes())

    print(f"Fichier .eml généré : {eml_path}")

def generate_with_mail(intervenant, societe, contact, duree, date_deb, date_fin, obj, contenu, mission):
    mail_intervenant = get_mail_intervenant(intervenant)
    pdf_path = generate_document(intervenant, societe, contact, duree, date_deb, date_fin, obj, contenu, mission, mail_intervenant)
    mail_contact = get_mail_contact(contact)
    prepare_outlook_email(mail_contact, mail_intervenant, pdf_path, societe)
    return pdf_path

def interface():
    clients = get_clients()
    intervenants = get_intervenants()

    with gr.Blocks() as demo:
        with gr.Tab("Générer Bon d'Intervention"):
            intervenant = gr.Dropdown(label="Intervenant", choices=intervenants)
            societe = gr.Dropdown(label="Société", choices=clients)
            contact = gr.Dropdown(label="Contact", choices=[])
            duree = gr.Textbox(label="Durée")
            date_deb = gr.Textbox(label="Date début")
            date_fin = gr.Textbox(label="Date fin")
            obj = gr.Textbox(label="Objectif")
            contenu = gr.Textbox(label="Contenu")
            mission = gr.Textbox(label="Numéro de mission")
            fichier_pdf = gr.File(label="Bon d'intervention (PDF)")

            def update_contacts(soc):
                return gr.update(choices=get_contacts(soc))
            societe.change(update_contacts, inputs=societe, outputs=contact)

            gr.Button("Générer PDF").click(generate_with_mail, 
                inputs=[intervenant, societe, contact, duree, date_deb, date_fin, obj, contenu, mission],
                outputs=fichier_pdf)

        with gr.Tab("Ajouter Client"):
            new_client = gr.Textbox(label="Société")
            msg_client = gr.Textbox(label="Message", interactive=False)
            gr.Button("Ajouter Client").click(add_client, inputs=new_client, outputs=msg_client)

        with gr.Tab("Ajouter Contact"):
            societe_contact = gr.Dropdown(label="Société", choices=clients)
            nom = gr.Textbox(label="Nom")
            prenom = gr.Textbox(label="Prénom")
            mail = gr.Textbox(label="Email")
            tel = gr.Textbox(label="Téléphone")
            msg_contact = gr.Textbox(label="Message", interactive=False)
            gr.Button("Ajouter Contact").click(add_contact, 
                inputs=[societe_contact, nom, prenom, mail, tel], outputs=msg_contact)

        with gr.Tab("Tableau de Bord"):
            data = get_bon_intervention()
            gr.DataFrame(data, headers=[
                "ID", "Intervenant", "Société", "Nom Contact", "Email Contact",
                "Durée", "Date Début", "Date Fin", "Objectif", "Contenu",
                "Numéro de Mission", "Date Création"
            ])

    return demo

if __name__ == "__main__":
    demo = interface()
    demo.launch(server_name="0.0.0.0", server_port=7860)
